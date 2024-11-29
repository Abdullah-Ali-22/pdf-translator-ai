# backend.py

import os
import re
from typing import Tuple, Set
import numpy as np
from dotenv import load_dotenv
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import ContentFormat
from tensorflow.keras.applications.vgg16 import VGG16, preprocess_input
from tensorflow.keras.preprocessing.image import img_to_array
from tensorflow.keras.models import Model

# Import the translate_text function from translator.py
from translator import translate_text

# Import the extract_job_details function and JobDetailsSchema from table_generate_agent.py
from table_generate_agent import extract_job_details, JobDetailsSchema

# Load environment variables
load_dotenv()

# Azure service credentials
DOCUMENT_INT_ENDPOINT = os.getenv("AZURE_DOCUMENTINT_ENDPOINT")
DOCUMENT_INT_KEY = os.getenv("AZURE_DOCUMENTINT_KEY")

# Load the VGG16 model pretrained on ImageNet
BASE_MODEL = VGG16(weights='imagenet')
MODEL = Model(inputs=BASE_MODEL.input, outputs=BASE_MODEL.get_layer('fc1').output)


def ensure_folder_exists(folder_path: str):
    """Create the folder if it does not exist."""
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Created folder at {folder_path}")
    else:
        print(f"Folder already exists at {folder_path}")


def preprocess_image(image: Image.Image, target_size: Tuple[int, int] = (224, 224)) -> np.ndarray:
    """Preprocess the image for feature extraction."""
    image = image.resize(target_size)
    image_array = img_to_array(image)
    image_array = np.expand_dims(image_array, axis=0)
    image_array = preprocess_input(image_array)
    return image_array


def extract_features(image: Image.Image) -> np.ndarray:
    """Extract features from a PIL image."""
    image_array = preprocess_image(image)
    features = MODEL.predict(image_array)
    return features.flatten()


def is_logo(image: Image.Image, logo_image_features: np.ndarray, similarity_threshold: float = 0.9) -> bool:
    """Determine if an image is a logo based on similarity to the logo image."""
    image_features = extract_features(image)
    similarity = np.dot(image_features, logo_image_features) / (
        np.linalg.norm(image_features) * np.linalg.norm(logo_image_features)
    )
    return similarity > similarity_threshold


def crop_image_from_pdf_page(pdf_path: str, page_number: int, bounding_box: Tuple[float, float, float, float]) -> Image.Image:
    """
    Crop a region from a given page in a PDF and return it as an image.

    Args:
        pdf_path (str): Path to the PDF file.
        page_number (int): Page number (0-indexed).
        bounding_box (tuple): Bounding box coordinates (x0, y0, x1, y1).

    Returns:
        PIL.Image.Image: Cropped image.
    """
    with fitz.open(pdf_path) as doc:
        page = doc.load_page(page_number)
        rect = fitz.Rect(
            bounding_box[0] * 72,
            bounding_box[1] * 72,
            bounding_box[2] * 72,
            bounding_box[3] * 72,
        )
        pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72), clip=rect)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img


def analyze_layout(
    input_file_path: str,
    images_output_folder: str,
    logo_output_folder: str,
    logo_image_features: np.ndarray,
    similarity_threshold: float = 0.8,
) -> Tuple[str, Set[int]]:
    """
    Analyze the layout of a document and extract figures.

    Args:
        input_file_path (str): Path to the input PDF file.
        images_output_folder (str): Folder to save non-logo images.
        logo_output_folder (str): Folder to save logo images.
        logo_image_features (np.ndarray): Features of the logo image.
        similarity_threshold (float): Threshold to determine if an image is a logo.

    Returns:
        Tuple[str, Set[int]]: Markdown content extracted from the document and set of logo figure indices.
    """
    ensure_folder_exists(images_output_folder)
    ensure_folder_exists(logo_output_folder)

    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=DOCUMENT_INT_ENDPOINT,
        credential=AzureKeyCredential(DOCUMENT_INT_KEY),
    )

    with open(input_file_path, "rb") as f:
        poller = document_intelligence_client.begin_analyze_document(
            "prebuilt-layout",
            analyze_request=f,
            content_type="application/octet-stream",
            output_content_format=ContentFormat.MARKDOWN,
        )

    result = poller.result()
    md_content = result.content

    logo_fig_indices = set()  # To track which figures are logos
    non_logo_idx = 0  # Counter for non-logo images

    for idx, figure in enumerate(result.figures):
        for region in figure.bounding_regions:
            # Extract x and y coordinates from the flat list
            polygon = region.polygon  # This is a list of floats

            xs = polygon[::2]  # x coordinates at even indices
            ys = polygon[1::2]  # y coordinates at odd indices

            # Compute bounding box coordinates
            x0, y0 = min(xs), min(ys)
            x1, y1 = max(xs), max(ys)
            bounding_box = (x0, y0, x1, y1)

            # Proceed with cropping and saving the image
            cropped_image = crop_image_from_pdf_page(
                input_file_path, region.page_number - 1, bounding_box
            )

            # Determine if the cropped image is a logo
            if is_logo(cropped_image, logo_image_features, similarity_threshold):
                # Save logo image
                image_filename = f"{idx}.png"
                image_path = os.path.join(logo_output_folder, image_filename)
                cropped_image.save(image_path)
                logo_fig_indices.add(idx)  # Mark this figure as a logo
                print(f"Logo {idx} cropped and saved as {image_path}")
            else:
                # Save non-logo image with sequential naming
                image_filename = f"{non_logo_idx}.png"
                image_path = os.path.join(images_output_folder, image_filename)
                cropped_image.save(image_path)
                print(f"Figure {idx} cropped and saved as {image_path}")
                non_logo_idx += 1  # Increment non-logo image counter

    return md_content, logo_fig_indices


def clean_md_content(md_content: str) -> str:
    """
    Clean up markdown content by removing unwanted lines and tags.

    Args:
        md_content (str): Original markdown content.

    Returns:
        str: Cleaned markdown content.
    """
    lines = md_content.splitlines()
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        # Remove HTML comments
        if re.match(r'<!--.*-->', line):
            continue
        # Remove empty lines
        if not line:
            continue
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)


def save_translated_word(
    job_details: JobDetailsSchema,
    translated_content: str,
    word_output_path: str,
    images_output_folder: str
) -> None:
    """
    Save the translated content along with job details and images to a Word document.

    Args:
        job_details (JobDetailsSchema): Extracted job details to include in the table.
        translated_content (str): User-edited translated content.
        word_output_path (str): Path to save the translated Word document.
        images_output_folder (str): Folder containing non-logo images.
    """
    # Create a new Word document
    translated_doc = Document()

    # Insert the job details table at the beginning
    if job_details:
        table = translated_doc.add_table(rows=0, cols=2)
        table.style = 'Light List Accent 1'  # Choose a style as needed

        # Helper function to add a row to the table
        def add_table_row(table, key, value):
            row = table.add_row().cells
            row[0].text = key.replace('_', ' ').capitalize()
            row[1].text = value if value is not None else "N/A"

        # Add rows for each field in JobDetailsSchema
        for field, value in job_details.dict().items():
            if field == 'job_description' and value:
                for sub_field, sub_value in value.dict().items():
                    formatted_key = sub_field.replace('_', ' ').capitalize()
                    formatted_value = ', '.join(sub_value) if isinstance(sub_value, list) else sub_value
                    add_table_row(table, formatted_key, formatted_value)
            else:
                formatted_key = field.replace('_', ' ').capitalize()
                add_table_row(table, formatted_key, value if value is not None else "N/A")

        # Add a paragraph after the table for spacing
        translated_doc.add_paragraph("\n")

    # Add the user-edited translated content
    if translated_content:
        # Split the content into paragraphs based on double newlines
        paragraphs = translated_content.split('\n\n')
        for para in paragraphs:
            if para.startswith('#'):
                # Handle headings
                heading_level = len(para) - len(para.lstrip('#'))
                heading_text = para.lstrip('#').strip()
                if heading_level <= 6:
                    translated_doc.add_heading(heading_text, level=heading_level)
                else:
                    translated_doc.add_paragraph(heading_text)
            else:
                translated_doc.add_paragraph(para)

    # Insert images at the end of the document
    image_files = sorted([
        f for f in os.listdir(images_output_folder)
        if os.path.isfile(os.path.join(images_output_folder, f)) and f.lower().endswith(('.png', '.jpg', '.jpeg'))
    ], key=lambda x: int(os.path.splitext(x)[0]))

    for image_filename in image_files:
        image_path = os.path.join(images_output_folder, image_filename)
        if os.path.exists(image_path):
            translated_doc.add_picture(image_path, width=Inches(4))  # Adjust size as needed
            last_paragraph = translated_doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"Inserted image {image_filename} into the translated document.")
        else:
            print(f"Image not found: {image_path}")

    # Save the translated Word document
    translated_doc.save(word_output_path)
    print(f"Translated Word document saved at {word_output_path}")