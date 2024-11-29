# backend.py

import os
import re
import uuid
import base64
import mimetypes
import numpy as np
from dotenv import load_dotenv
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import ContentFormat
from tensorflow.keras.applications.vgg16 import VGG16, preprocess_input
from tensorflow.keras.preprocessing.image import img_to_array
from tensorflow.keras.models import Model

# Import the translate_text function from translator.py
from translator import translate_text

# Import the extract_job_details function and JobDetailsSchema from table_generate_agent.py
from table_generate_agent import extract_job_details, JobDetailsSchema

# Load environment variables
load_dotenv()

# Azure service credentials
DOCUMENT_INT_ENDPOINT = os.getenv("AZURE_DOCUMENTINT_ENDPOINT")
DOCUMENT_INT_KEY = os.getenv("AZURE_DOCUMENTINT_KEY")

# Load the VGG16 model pretrained on ImageNet
BASE_MODEL = VGG16(weights='imagenet')
MODEL = Model(inputs=BASE_MODEL.input, outputs=BASE_MODEL.get_layer('fc1').output)

def ensure_folder_exists(folder_path):
    """Create the folder if it does not exist."""
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Created folder at {folder_path}")
    else:
        print(f"Folder already exists at {folder_path}")

def preprocess_image(image, target_size=(224, 224)):
    """Preprocess the image for feature extraction."""
    image = image.resize(target_size)
    image_array = img_to_array(image)
    image_array = np.expand_dims(image_array, axis=0)
    image_array = preprocess_input(image_array)
    return image_array

def extract_features(image):
    """Extract features from a PIL image."""
    image_array = preprocess_image(image)
    features = MODEL.predict(image_array)
    return features.flatten()

def is_logo(image, logo_image_features, similarity_threshold=0.9):
    """Determine if an image is a logo based on similarity to the logo image."""
    image_features = extract_features(image)
    similarity = np.dot(image_features, logo_image_features) / (
        np.linalg.norm(image_features) * np.linalg.norm(logo_image_features)
    )
    return similarity > similarity_threshold

def crop_image_from_pdf_page(pdf_path, page_number, bounding_box):
    """
    Crop a region from a given page in a PDF and return it as an image.

    Args:
        pdf_path (str): Path to the PDF file.
        page_number (int): Page number (0-indexed).
        bounding_box (tuple): Bounding box coordinates (x0, y0, x1, y1).

    Returns:
        PIL.Image.Image: Cropped image.
    """
    with fitz.open(pdf_path) as doc:
        page = doc.load_page(page_number)
        rect = fitz.Rect(
            bounding_box[0] * 72,
            bounding_box[1] * 72,
            bounding_box[2] * 72,
            bounding_box[3] * 72,
        )
        pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72), clip=rect)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img

def analyze_layout(
    input_file_path,
    images_output_folder,
    logo_output_folder,
    logo_image_features,
    similarity_threshold=0.8,
):
    """
    Analyze the layout of a document and extract figures.

    Args:
        input_file_path (str): Path to the input PDF file.
        images_output_folder (str): Folder to save non-logo images.
        logo_output_folder (str): Folder to save logo images.
        logo_image_features (np.ndarray): Features of the logo image.
        similarity_threshold (float): Threshold to determine if an image is a logo.

    Returns:
        str: Markdown content extracted from the document.
        set: Set of figure indices identified as logos.
    """
    ensure_folder_exists(images_output_folder)
    ensure_folder_exists(logo_output_folder)

    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=DOCUMENT_INT_ENDPOINT,
        credential=AzureKeyCredential(DOCUMENT_INT_KEY),
    )

    with open(input_file_path, "rb") as f:
        poller = document_intelligence_client.begin_analyze_document(
            "prebuilt-layout",
            analyze_request=f,
            content_type="application/octet-stream",
            output_content_format=ContentFormat.MARKDOWN,
        )

    result = poller.result()
    md_content = result.content

    logo_fig_indices = set()       # To track which figures are logos
    non_logo_idx = 0               # Counter for non-logo images

    for idx, figure in enumerate(result.figures):
        for region in figure.bounding_regions:
            # Extract x and y coordinates from the flat list
            polygon = region.polygon  # This is a list of floats

            xs = polygon[::2]  # x coordinates at even indices
            ys = polygon[1::2]  # y coordinates at odd indices

            # Compute bounding box coordinates
            x0, y0 = min(xs), min(ys)
            x1, y1 = max(xs), max(ys)
            bounding_box = (x0, y0, x1, y1)

            # Proceed with cropping and saving the image
            cropped_image = crop_image_from_pdf_page(
                input_file_path, region.page_number - 1, bounding_box
            )

            # Determine if the cropped image is a logo
            if is_logo(cropped_image, logo_image_features, similarity_threshold):
                # Save logo image
                image_filename = f"{idx}.png"
                image_path = os.path.join(logo_output_folder, image_filename)
                cropped_image.save(image_path)
                logo_fig_indices.add(idx)  # Mark this figure as a logo
                print(f"Logo {idx} cropped and saved as {image_path}")
            else:
                # Save non-logo image with sequential naming
                image_filename = f"{non_logo_idx}.png"
                image_path = os.path.join(images_output_folder, image_filename)
                cropped_image.save(image_path)
                print(f"Figure {idx} cropped and saved as {image_path}")
                non_logo_idx += 1  # Increment non-logo image counter

    return md_content, logo_fig_indices

def clean_md_content(md_content):
    """
    Clean up markdown content by removing unwanted lines and tags.

    Args:
        md_content (str): Original markdown content.

    Returns:
        str: Cleaned markdown content.
    """
    lines = md_content.splitlines()
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        # Remove HTML comments
        if re.match(r'<!--.*-->', line):
            continue
        # Remove empty lines
        if not line:
            continue
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)

def save_figures_to_word_with_position(md_content, images_output_folder, word_output_path, logo_fig_indices):
    """
    Save figures in the Word document in the same order as they appear in the document,
    excluding logo images.

    Args:
        md_content (str): Markdown content extracted from the document.
        images_output_folder (str): Folder containing the images.
        word_output_path (str): Path to save the Word document.
        logo_fig_indices (set): Set of figure indices identified as logos.
    """
    # Clean up the markdown content
    md_content = clean_md_content(md_content)

    # Create a new Word document
    doc = Document()

    # Split markdown content into lines
    lines = md_content.splitlines()

    figure_counter = 0       # To keep track of figure indices in markdown
    non_logo_insert_counter = 0  # To track images in images_output_folder
    in_figure = False        # Flag to indicate if we're inside a figure block

    for line in lines:
        line = line.strip()

        # Handle figure start tag
        if line.lower() == '<figure>':
            in_figure = True
            figure_counter += 1  # Increment figure_counter as we're entering a new figure
            continue

        # Handle figure end tag
        elif line.lower() == '</figure>':
            in_figure = False

            # Check if the current figure is a logo
            current_fig_idx = figure_counter - 1  # Zero-based index
            if current_fig_idx in logo_fig_indices:
                print(f"Skipping logo figure {current_fig_idx} from the document.")
                continue  # Skip inserting this image

            # Insert the figure image
            image_filename = f"{non_logo_insert_counter}.png"
            image_path = os.path.join(images_output_folder, image_filename)

            if os.path.exists(image_path):
                # Add image to Word document
                doc.add_picture(image_path, width=Inches(4))  # Adjust size as needed
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"Inserted image {image_filename} into the document.")
                non_logo_insert_counter += 1  # Increment after inserting a non-logo image
            else:
                print(f"Image not found: {image_path}")

            continue

        # Handle content inside figure (if any)
        if in_figure:
            continue  # Skip content inside figure tags if not needed

        # Handle headings
        if line.startswith('#'):
            heading_level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()
            if heading_level <= 6:
                doc.add_heading(heading_text, level=heading_level)
            else:
                doc.add_paragraph(heading_text)
        else:
            # Add the line as a paragraph
            doc.add_paragraph(line)

    # Save the Word document
    doc.save(word_output_path)
    print(f"Word document saved at {word_output_path}")

def save_translated_word(md_content, images_output_folder, word_output_path, logo_output_folder, logo_fig_indices, job_details: JobDetailsSchema):
    """
    Save a translated version of the document to a Word file, including a job details table.

    Args:
        md_content (str): Original markdown content.
        images_output_folder (str): Folder containing non-logo images.
        word_output_path (str): Path to save the translated Word document.
        logo_output_folder (str): Folder containing logo images.
        logo_fig_indices (set): Set of figure indices identified as logos.
        job_details (JobDetailsSchema): Extracted job details to include in the table.
    """
    # Clean up the markdown content
    md_content = clean_md_content(md_content)

    # Split the markdown content into lines
    lines = md_content.splitlines()

    # Initialize variables
    translated_doc = Document()
    figure_counter = 0       # To keep track of figure indices in markdown
    non_logo_insert_counter = 0  # To track images in images_output_folder
    in_figure = False        # Flag to indicate if we're inside a figure block
    text_buffer = []         # Buffer to collect text lines

    # Insert the job details table at the beginning
    if job_details:
        table = translated_doc.add_table(rows=0, cols=2)
        table.style = 'Light List Accent 1'  # Choose a style as needed

        # Helper function to add a row to the table
        def add_table_row(table, key, value):
            row = table.add_row().cells
            row[0].text = key.replace('_', ' ').capitalize()
            row[1].text = value if value is not None else "N/A"

        # Add rows for each field in JobDetailsSchema
        add_table_row(table, 'role_position', job_details.role_position)
        add_table_row(table, 'location', job_details.location)
        add_table_row(table, 'number_of_fte', job_details.number_of_fte)
        add_table_row(table, 'rgs_id', job_details.rgs_id)
        add_table_row(table, 'remote_onsite', job_details.remote_onsite)
        add_table_row(table, 'onsite_frequency_week', job_details.onsite_frequency_week)
        add_table_row(table, 'project_duration', job_details.project_duration)
        add_table_row(table, 'working_hours_per_day', job_details.working_hours_per_day)
        add_table_row(table, 'contract_mode', job_details.contract_mode)
        add_table_row(table, 'daily_rate', job_details.daily_rate)
        add_table_row(table, 'language_proficiency', job_details.language_proficiency)
        add_table_row(table, 'start_date_of_engagement', job_details.start_date_of_engagement)
        add_table_row(table, 'experience_required', job_details.experience_required)

        # If JobDetailsSchema has nested JobDescriptionSchema, handle it
        if job_details.job_description:
            job_desc = job_details.job_description
            add_table_row(table, 'must', ', '.join(job_desc.must) if job_desc.must else "N/A")
            add_table_row(table, 'target', ', '.join(job_desc.target) if job_desc.target else "N/A")

        # Add a paragraph after the table for spacing
        translated_doc.add_paragraph("\n")

    for line in lines:
        line = line.strip()

        # Handle figure start tag
        if line.lower() == '<figure>':
            in_figure = True
            figure_counter += 1  # Increment figure_counter as we're entering a new figure

            # If there's accumulated text, translate and insert it
            if text_buffer:
                text_to_translate = '\n'.join(text_buffer)
                translations = translate_text(text_to_translate)
                # Assuming the translation returns a list of translations
                # Here, we take the first translation (adjust if multiple translations are returned)
                if translations and len(translations) > 0:
                    translated_header = translations[0].header or ""
                    translated_content = translations[0].content or ""
                    # Insert translated header if exists
                    if translated_header:
                        translated_doc.add_heading(translated_header, level=1)
                    # Insert translated content
                    if translated_content:
                        translated_doc.add_paragraph(translated_content)
                text_buffer = []  # Clear the buffer

            continue

        # Handle figure end tag
        elif line.lower() == '</figure>':
            in_figure = False

            # Check if the current figure is a logo
            current_fig_idx = figure_counter - 1  # Zero-based index
            if current_fig_idx in logo_fig_indices:
                print(f"Skipping logo figure {current_fig_idx} from the translated document.")
                continue  # Skip inserting this image

            # Insert the figure image
            image_filename = f"{non_logo_insert_counter}.png"
            image_path = os.path.join(images_output_folder, image_filename)

            if os.path.exists(image_path):
                # Add image to translated Word document
                translated_doc.add_picture(image_path, width=Inches(4))  # Adjust size as needed
                last_paragraph = translated_doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                print(f"Inserted image {image_filename} into the translated document.")
                non_logo_insert_counter += 1  # Increment after inserting a non-logo image
            else:
                print(f"Image not found: {image_path}")

            continue

        # Handle content inside figure (if any)
        if in_figure:
            continue  # Skip content inside figure tags if not needed

        # Collect text lines
        text_buffer.append(line)

    # After processing all lines, check if there's remaining text to translate
    if text_buffer:
        text_to_translate = '\n'.join(text_buffer)
        translations = translate_text(text_to_translate)
        if translations and len(translations) > 0:
            translated_header = translations[0].header or ""
            translated_content = translations[0].content or ""
            # Insert translated header if exists
            if translated_header:
                translated_doc.add_heading(translated_header, level=1)
            # Insert translated content
            if translated_content:
                translated_doc.add_paragraph(translated_content)

    # Save the translated Word document
    translated_doc.save(word_output_path)
    print(f"Translated Word document saved at {word_output_path}")